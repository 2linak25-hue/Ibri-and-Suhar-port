"""
Tables Module
Functions for creating various tables in the Word document
"""

from docx.shared import Pt, Inches, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_cell_shading(cell, color_hex):
    """Set cell background color"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color_hex)
    cell._tc.get_or_add_tcPr().append(shading_elm)


def set_table_borders(table):
    """Set table borders"""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:color'), 'CCCCCC')
        tblBorders.append(border)
    
    tblPr.append(tblBorders)
    if tbl.tblPr is None:
        tbl.insert(0, tblPr)


def create_table(document, headers, rows, caption=None):
    """
    Create a formatted table with headers and data rows
    
    Args:
        document: Word document object
        headers: List of header strings
        rows: List of row data (each row is a list)
        caption: Optional table caption
    """
    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table)
    
    # Header row
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        cell = header_cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(10)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, 'E8E8E8')
    
    # Data rows
    for row_data in rows:
        row_cells = table.add_row().cells
        for i, value in enumerate(row_data):
            cell = row_cells[i]
            cell.text = str(value)
            cell.paragraphs[0].runs[0].font.size = Pt(9)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add caption if provided
    if caption:
        p = document.add_paragraph()
        run = p.add_run(caption)
        run.font.size = Pt(10)
        run.font.italic = True
        run.font.color.rgb = RGBColor(100, 100, 100)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    document.add_paragraph()  # Spacing
    return table


def create_results_table(document, title, headers, rows, highlight_best=False):
    """
    Create a results table with optional highlighting
    
    Args:
        document: Word document object
        title: Table title
        headers: List of header strings
        rows: List of row data
        highlight_best: Whether to highlight the best row
    """
    # Add title
    p = document.add_paragraph()
    run = p.add_run(title)
    run.font.bold = True
    run.font.size = Pt(11)
    
    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table)
    
    # Header row
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        cell = header_cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(9)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, '2C3E50')
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    
    # Data rows
    for idx, row_data in enumerate(rows):
        row_cells = table.add_row().cells
        is_best = highlight_best and row_data[0] == 'XGBoost'
        
        for i, value in enumerate(row_data):
            cell = row_cells[i]
            cell.text = str(value)
            para = cell.paragraphs[0]
            para.runs[0].font.size = Pt(9)
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            if is_best:
                para.runs[0].font.bold = True
                set_cell_shading(cell, 'E8F6E8')
            elif idx % 2 == 0:
                set_cell_shading(cell, 'F8F8F8')
    
    document.add_paragraph()
    return table


def create_descriptive_stats_table(document):
    """Create the descriptive statistics table"""
    headers = ['Feature', 'Mean', 'Median', 'Std', 'Min', 'Max', 'Skewness']
    rows = [
        ['Distance_m', '112.58', '50.0', '119.61', '0.0', '390.0', '0.92'],
        ['Temp_C', '30.37', '29.4', '1.47', '29.0', '33.1', '0.83'],
        ['Humidity_Pct', '35.22', '36.3', '4.06', '30.4', '40.8', '0.26'],
        ['Time_Hour', '12.36', '15.0', '3.09', '8.17', '15.0', '-0.40'],
        ['E_ICNIRP', '10.69', '11.56', '5.87', '0.17', '21.55', '-0.17'],
        ['H_ICNIRP', '3.47', '3.54', '1.51', '0.55', '6.15', '-0.19'],
    ]
    return create_table(document, headers, rows, 'Table 1: Descriptive Statistics Summary')


def create_anova_table(document):
    """Create the ANOVA results table"""
    headers = ['Feature', 'Target', 'F-Statistic', 'p-value', 'Eta²', 'Significant']
    rows = [
        ['City', 'E_ICNIRP', '1.36', '0.247', '0.021', 'No'],
        ['City', 'H_ICNIRP', '14.01', '0.0004', '0.180', 'Yes'],
        ['Profile_Type', 'E_ICNIRP', '1.71', '0.195', '0.026', 'No'],
        ['Profile_Type', 'H_ICNIRP', '10.23', '0.002', '0.138', 'Yes'],
        ['Circuit', 'E_ICNIRP', '0.96', '0.387', '0.030', 'No'],
        ['Circuit', 'H_ICNIRP', '6.97', '0.002', '0.181', 'Yes'],
    ]
    return create_table(document, headers, rows, 'Table 2: ANOVA Results')


def create_vif_table(document):
    """Create the VIF multicollinearity table"""
    headers = ['Feature', 'VIF', 'Status']
    rows = [
        ['Distance_m', '1.83', 'OK (<5)'],
        ['Circuit', '5.21', 'MODERATE (5-10)'],
        ['City', '∞', 'HIGH (>10)'],
        ['Profile_Type', '∞', 'HIGH (>10)'],
        ['Time_Hour', '∞', 'HIGH (>10)'],
        ['Temp_C', '∞', 'HIGH (>10)'],
        ['Humidity_Pct', '∞', 'HIGH (>10)'],
    ]
    return create_table(document, headers, rows, 'Table 3: Variance Inflation Factor (VIF) Results')


def create_model_results_E_table(document):
    """Create model results table for E_ICNIRP"""
    headers = ['Model', 'Train R²', 'Test R²', 'Test RMSE', 'Test MAE', 'CV R² (Mean±Std)']
    rows = [
        ['SVR', '0.471', '-0.112', '5.70', '4.50', '0.047 ± 0.271'],
        ['Random Forest', '0.684', '-0.067', '5.58', '4.46', '0.259 ± 0.163'],
        ['XGBoost', '0.722', '0.269', '4.62', '3.52', '0.173 ± 0.265'],
        ['Neural Network', '0.340', '-0.550', '6.73', '5.32', '-0.332 ± 0.452'],
    ]
    return create_results_table(document, 'E_ICNIRP Target Performance', headers, rows, highlight_best=True)


def create_model_results_H_table(document):
    """Create model results table for H_ICNIRP"""
    headers = ['Model', 'Train R²', 'Test R²', 'Test RMSE', 'Test MAE', 'CV R² (Mean±Std)']
    rows = [
        ['SVR', '0.681', '-0.271', '1.17', '0.78', '0.204 ± 0.810'],
        ['Random Forest', '0.760', '0.401', '0.80', '0.67', '0.217 ± 0.587'],
        ['XGBoost', '0.716', '0.535', '0.71', '0.56', '0.247 ± 0.634'],
        ['Neural Network', '0.079', '-0.898', '1.43', '1.17', '-0.099 ± 0.559'],
    ]
    return create_results_table(document, 'H_ICNIRP Target Performance', headers, rows, highlight_best=True)


def create_feature_importance_table(document):
    """Create feature importance ranking table"""
    headers = ['Rank', 'Feature', 'Avg Importance', 'Interpretation']
    rows = [
        ['1', 'Dist_Temp_Interaction', '0.841', 'Distance-Temperature interaction'],
        ['2', 'Temp_C', '0.591', 'Temperature influence'],
        ['3', 'Distance_m', '0.561', 'Inverse square law'],
        ['4', 'Distance_x_Humidity', '0.543', 'Distance-Humidity interaction'],
        ['5', 'Distance_Squared', '0.403', 'Non-linear distance effect'],
        ['6', 'Distance_Inverse', '0.396', 'Inverse distance'],
        ['7', 'Dist_Hum_Interaction', '0.365', 'Environmental-spatial'],
        ['8', 'Humidity_Pct', '0.244', 'Humidity impact'],
        ['9', 'Circuit', '0.234', 'Hardware configuration'],
        ['10', 'Profile_Type', '0.217', 'Measurement profile'],
    ]
    return create_table(document, headers, rows, 'Table 4: Top 10 Feature Importance Rankings')


def create_normality_test_table(document):
    """Create normality test results table"""
    headers = ['Feature', 'Shapiro-Wilk p-value', 'Normal?', 'Anderson-Darling']
    rows = [
        ['E_ICNIRP', '0.071', 'Yes', 'Pass'],
        ['H_ICNIRP', '0.045', 'No', 'Pass'],
        ['Distance_m', '3.07e-07', 'No', 'Fail'],
        ['Temp_C', '2.00e-08', 'No', 'Fail'],
        ['Humidity_Pct', '1.33e-07', 'No', 'Fail'],
    ]
    return create_table(document, headers, rows, 'Table 5: Normality Test Results')


def create_chi_square_table(document):
    """Create chi-square test results table"""
    headers = ['Variable 1', 'Variable 2', 'χ²', 'p-value', "Cramér's V", 'Significant']
    rows = [
        ['City', 'Profile_Type', '0.39', '0.535', '0.076', 'No'],
        ['City', 'Circuit', '66.0', '4.66e-15', '1.000', 'Yes'],
        ['Profile_Type', 'Circuit', '0.79', '0.674', '0.109', 'No'],
    ]
    return create_table(document, headers, rows, 'Table 6: Chi-Square Test Results')


def create_effect_size_table(document):
    """Create effect size table"""
    headers = ['Feature', 'E_ICNIRP Effect', 'H_ICNIRP Effect']
    rows = [
        ['City', 'Small (0.021)', 'Large (0.180)'],
        ['Profile_Type', 'Small (0.026)', 'Medium (0.138)'],
        ['Circuit', 'Small (0.030)', 'Large (0.181)'],
    ]
    return create_table(document, headers, rows, 'Table 7: Effect Size Analysis (Eta-Squared)')


def create_metrics_explanation_table(document):
    """Create evaluation metrics explanation table"""
    headers = ['Metric', 'Formula', 'Interpretation']
    rows = [
        ['RMSE', '√(Σ(y-ŷ)²/n)', 'Penalizes large errors'],
        ['MAE', 'Σ|y-ŷ|/n', 'Average absolute error'],
        ['R²', '1 - SS_res/SS_tot', 'Variance explained (0-1)'],
        ['MAPE', '100/n × Σ|(y-ŷ)/y|', 'Percentage error'],
    ]
    return create_table(document, headers, rows, 'Table 8: Evaluation Metrics')
