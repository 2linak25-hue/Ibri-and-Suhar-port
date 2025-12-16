"""
Images Module
Functions for adding images to the Word document
"""

import os
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def add_image(document, image_path, caption, width=5.5):
    """
    Add an image with caption to the document
    
    Args:
        document: Word document object
        image_path: Full path to the image file
        caption: Caption text for the image
        width: Image width in inches (default 5.5)
    
    Returns:
        bool: True if image was added, False if not found
    """
    if os.path.exists(image_path):
        # Add centered paragraph for image
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(image_path, width=Inches(width))
        
        # Add caption
        cap_p = document.add_paragraph()
        cap_run = cap_p.add_run(caption)
        cap_run.font.size = Pt(10)
        cap_run.font.italic = True
        cap_run.font.color.rgb = RGBColor(100, 100, 100)
        cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap_p.paragraph_format.space_after = Pt(12)
        
        return True
    else:
        # Add placeholder text if image not found
        p = document.add_paragraph()
        run = p.add_run(f'[Image not found: {os.path.basename(image_path)}]')
        run.font.italic = True
        run.font.color.rgb = RGBColor(200, 100, 100)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        return False


def add_all_images(document, plots_dir):
    """
    Add all analysis images to the document
    
    Args:
        document: Word document object
        plots_dir: Directory containing the plot images
    
    Returns:
        dict: Statistics about images added
    """
    images = {
        # Data Quality Section
        'data_quality': [
            ('01_missing_values_heatmap.png', 'Figure 1: Missing Values Heatmap - Visualization of data completeness across all features'),
        ],
        # Correlation Analysis
        'correlation': [
            ('02_correlation_heatmap.png', 'Figure 2: Correlation Heatmap - Pearson correlation coefficients between all numerical features'),
        ],
        # VIF Analysis
        'vif': [
            ('11_vif_multicollinearity.png', 'Figure 3: Variance Inflation Factor Analysis - Multicollinearity assessment'),
        ],
        # Data Exploration
        'exploration': [
            ('03_target_distribution.png', 'Figure 4: Distribution of Target Variables (E_ICNIRP and H_ICNIRP)'),
            ('04_boxplots_numerical.png', 'Figure 5: Box Plots for Numerical Features - Outlier detection'),
        ],
        # Feature Importance
        'features': [
            ('05_feature_importance.png', 'Figure 6: Aggregated Feature Importance Rankings from Tree-based Models'),
        ],
        # Model Comparison
        'comparison': [
            ('06_model_comparison_r2.png', 'Figure 7: Model Comparison - Test R² Scores for all models'),
            ('07_model_comparison_rmse.png', 'Figure 8: Model Comparison - Test RMSE (Lower is Better)'),
            ('12_model_dashboard.png', 'Figure 9: Comprehensive Model Comparison Dashboard'),
            ('model_comparison_with_stacked_ensemble.png', 'Figure 10: Model Comparison Including Stacked Ensemble Framework'),
        ],
        # Prediction Analysis
        'predictions': [
            ('08_actual_vs_predicted_E_ICNIRP.png', 'Figure 11: Actual vs Predicted Values for E_ICNIRP'),
            ('09_actual_vs_predicted_H_ICNIRP.png', 'Figure 12: Actual vs Predicted Values for H_ICNIRP'),
            ('10_residual_plots.png', 'Figure 13: Residual Analysis for Best Models'),
            ('stacked_ensemble_performance.png', 'Figure 14: Stacked Ensemble Model Performance Analysis'),
        ],
    }
    
    stats = {'added': 0, 'not_found': 0}
    
    return images, stats


def add_section_images(document, plots_dir, section_key, images_dict):
    """
    Add images for a specific section
    
    Args:
        document: Word document object
        plots_dir: Directory containing plots
        section_key: Key from the images dictionary
        images_dict: Dictionary of images by section
    
    Returns:
        int: Number of images added
    """
    added = 0
    if section_key in images_dict:
        for filename, caption in images_dict[section_key]:
            image_path = os.path.join(plots_dir, filename)
            if add_image(document, image_path, caption):
                added += 1
    return added
