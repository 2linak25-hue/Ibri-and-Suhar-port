"""
Document Builder Module
Main class for building the Word document
"""

import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT

from .styles import apply_styles
from .tables import (
    create_descriptive_stats_table,
    create_anova_table,
    create_vif_table,
    create_model_results_E_table,
    create_model_results_H_table,
    create_feature_importance_table,
    create_normality_test_table,
    create_chi_square_table,
    create_effect_size_table,
    create_metrics_explanation_table
)
from .images import add_image, add_all_images
from .content import get_methodology_content, get_results_content, get_discussion_content


class DocumentBuilder:
    """Builder class for creating the EMF ML Analysis Word document"""
    
    def __init__(self, output_path, plots_dir=None):
        """
        Initialize the document builder
        
        Args:
            output_path: Path where the Word document will be saved
            plots_dir: Directory containing plot images
        """
        self.output_path = output_path
        self.plots_dir = plots_dir or os.path.join(os.path.dirname(output_path), 'plots')
        self.document = Document()
        self.images_dict, _ = add_all_images(None, None)  # Get image definitions
        
    def build(self):
        """Build the complete document"""
        # Apply styles
        apply_styles(self.document)
        
        # Build sections
        self._add_title_page()
        self._add_methodology_section()
        self._add_results_section()
        self._add_discussion_section()
        self._add_references()
        self._add_appendix()
        
        # Save document
        self.document.save(self.output_path)
        print(f"Document saved: {self.output_path}")
        
        return self.output_path
    
    def _add_title_page(self):
        """Add title page"""
        # Main title
        title = self.document.add_paragraph()
        title_run = title.add_run('Electromagnetic Field (EMF) Prediction Using Machine Learning')
        title_run.font.size = Pt(24)
        title_run.font.bold = True
        title_run.font.color.rgb = RGBColor(44, 62, 80)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Subtitle
        subtitle = self.document.add_paragraph()
        sub_run = subtitle.add_run('A Stacked Ensemble Approach')
        sub_run.font.size = Pt(18)
        sub_run.font.color.rgb = RGBColor(52, 73, 94)
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        self.document.add_paragraph()
        
        # Document description
        desc = self.document.add_paragraph()
        desc_run = desc.add_run('Comprehensive Methodology, Results, and Discussion')
        desc_run.font.size = Pt(14)
        desc_run.font.italic = True
        desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        self.document.add_paragraph()
        self.document.add_paragraph()
        
        # Project info
        info = self.document.add_paragraph()
        info_run = info.add_run('EMF ML Analysis Project\nIbri and Suhar Port Study\nDecember 2024')
        info_run.font.size = Pt(12)
        info.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Page break
        self.document.add_page_break()
    
    def _add_paragraph(self, text, style='Normal', bold=False):
        """Add a styled paragraph"""
        p = self.document.add_paragraph(style=style)
        run = p.add_run(text)
        if bold:
            run.font.bold = True
        return p
    
    def _add_methodology_section(self):
        """Add methodology section"""
        content = get_methodology_content()
        
        # Part I Header
        self.document.add_heading('PART I: METHODOLOGY', 0)
        
        # 1. Introduction
        self.document.add_heading('1. Introduction', 1)
        self.document.add_heading('1.1 Research Context', 2)
        self._add_paragraph(content['introduction']['research_context'])
        
        self.document.add_heading('1.2 Research Objectives', 2)
        self._add_paragraph(content['introduction']['objectives'])
        
        # 2. Data Collection
        self.document.add_heading('2. Data Collection and Description', 1)
        self.document.add_heading('2.1 Dataset Overview', 2)
        self._add_paragraph(content['data_collection']['overview'])
        
        self.document.add_heading('2.2 Feature Categories', 2)
        self._add_paragraph(content['data_collection']['features']['spatial'])
        self._add_paragraph(content['data_collection']['features']['environmental'])
        self._add_paragraph(content['data_collection']['features']['technical'])
        self._add_paragraph(content['data_collection']['features']['temporal'])
        
        # 3. Data Preprocessing
        self.document.add_heading('3. Data Preprocessing Pipeline', 1)
        self.document.add_heading('3.1 Data Quality Assessment', 2)
        
        # Add missing values heatmap
        self._add_section_image('data_quality', 0)
        
        self._add_paragraph(content['preprocessing']['quality_assessment'])
        
        self.document.add_heading('3.2 Feature Engineering', 2)
        self._add_paragraph(content['preprocessing']['feature_engineering'])
        
        self.document.add_heading('3.3 Dimensionality Reduction', 2)
        self._add_paragraph(content['preprocessing']['dimensionality'])
        
        # 4. Statistical Analysis
        self.document.add_heading('4. Statistical Analysis Framework', 1)
        
        self.document.add_heading('4.1 Correlation Analysis', 2)
        self._add_section_image('correlation', 0)
        self._add_paragraph(content['statistical_analysis']['correlation'])
        
        self.document.add_heading('4.2 Variance Inflation Factor (VIF)', 2)
        self._add_section_image('vif', 0)
        self._add_paragraph(content['statistical_analysis']['vif'])
        create_vif_table(self.document)
        
        self.document.add_heading('4.3 ANOVA Analysis', 2)
        self._add_paragraph(content['statistical_analysis']['anova'])
        
        self.document.add_heading('4.4 Normality Tests', 2)
        self._add_paragraph(content['statistical_analysis']['normality'])
        create_normality_test_table(self.document)
        
        # 5. Machine Learning Framework
        self.document.add_heading('5. Machine Learning Framework', 1)
        
        self.document.add_heading('5.1 Model Selection Rationale', 2)
        self._add_paragraph(content['ml_framework']['svr'])
        self._add_paragraph(content['ml_framework']['rf'])
        self._add_paragraph(content['ml_framework']['xgb'])
        self._add_paragraph(content['ml_framework']['nn'])
        
        self.document.add_heading('5.2 Stacked Ensemble Framework', 2)
        self._add_paragraph(content['ml_framework']['stacked_ensemble'])
        
        self.document.add_heading('5.3 Cross-Validation Strategy', 2)
        self._add_paragraph(content['cross_validation']['content'])
        
        # 6. Evaluation Metrics
        self.document.add_heading('6. Evaluation Metrics', 1)
        create_metrics_explanation_table(self.document)
        
        self.document.add_page_break()
    
    def _add_results_section(self):
        """Add results section"""
        content = get_results_content()
        
        # Part II Header
        self.document.add_heading('PART II: RESULTS', 0)
        
        # 7. Data Exploration Results
        self.document.add_heading('7. Data Exploration Results', 1)
        
        self.document.add_heading('7.1 Dataset Statistics', 2)
        # Add distribution plots
        self._add_section_image('exploration', 0)
        self._add_section_image('exploration', 1)
        
        self._add_paragraph(content['data_exploration']['stats'])
        create_descriptive_stats_table(self.document)
        
        self.document.add_heading('7.2 Correlation Analysis Findings', 2)
        self._add_paragraph(content['data_exploration']['correlation'])
        
        self.document.add_heading('7.3 ANOVA Results', 2)
        create_anova_table(self.document)
        create_effect_size_table(self.document)
        
        self.document.add_heading('7.4 Chi-Square Tests', 2)
        create_chi_square_table(self.document)
        
        # 8. Model Performance Results
        self.document.add_heading('8. Model Performance Results', 1)
        
        self.document.add_heading('8.1 Individual Model Performance', 2)
        create_model_results_E_table(self.document)
        create_model_results_H_table(self.document)
        
        self.document.add_heading('8.2 Stacked Ensemble Performance', 2)
        self._add_paragraph(content['stacked_ensemble']['content'])
        
        self.document.add_heading('8.3 Feature Importance Analysis', 2)
        self._add_section_image('features', 0)
        create_feature_importance_table(self.document)
        
        # 9. Visualizations
        self.document.add_heading('9. Visualizations', 1)
        
        self.document.add_heading('9.1 Model Comparison Dashboard', 2)
        for i in range(4):
            self._add_section_image('comparison', i)
        
        self.document.add_heading('9.2 Prediction Analysis', 2)
        for i in range(4):
            self._add_section_image('predictions', i)
        
        self.document.add_page_break()
    
    def _add_discussion_section(self):
        """Add discussion section"""
        content = get_discussion_content()
        
        # Part III Header
        self.document.add_heading('PART III: DISCUSSION', 0)
        
        # 10. Interpretation
        self.document.add_heading('10. Interpretation of Results', 1)
        self._add_paragraph(content['interpretation']['model_analysis'])
        self._add_paragraph(content['interpretation']['physical'])
        
        # 11. Limitations
        self.document.add_heading('11. Limitations', 1)
        self._add_paragraph(content['limitations']['content'])
        
        # 12. Future Work
        self.document.add_heading('12. Future Work', 1)
        self._add_paragraph(content['future_work']['content'])
        
        # 13. Conclusions
        self.document.add_heading('13. Conclusions', 1)
        self._add_paragraph(content['conclusions']['content'])
        
        self.document.add_page_break()
    
    def _add_references(self):
        """Add references section"""
        self.document.add_heading('References', 1)
        
        references = [
            '1. ICNIRP Guidelines for Limiting Exposure to Electromagnetic Fields (2020)',
            '2. Breiman, L. (1996). Stacked Regressions. Machine Learning, 24, 49-64.',
            '3. Chen, T., & Guestrin, C. (2016). XGBoost: A Scalable Tree Boosting System.',
            '4. Wolpert, D. H. (1992). Stacked Generalization. Neural Networks, 5(2), 241-259.',
        ]
        
        for ref in references:
            self._add_paragraph(ref)
    
    def _add_appendix(self):
        """Add appendix section"""
        self.document.add_heading('Appendix', 1)
        
        self.document.add_heading('A. Software and Libraries', 2)
        software = '''• Python 3.12
• scikit-learn 1.5.2
• XGBoost 3.0.3
• pandas 2.2.0
• numpy 1.26.2
• matplotlib/seaborn for visualization'''
        self._add_paragraph(software)
        
        self.document.add_heading('B. Model Artifacts', 2)
        artifacts = '''• Trained models: models/ directory
• Plots: outputs/plots/ directory
• Tables: outputs/tables/ directory'''
        self._add_paragraph(artifacts)
        
        self.document.add_heading('C. Reproducibility', 2)
        repro = '''• Random State: 42
• Cross-Validation: 5-fold
• Test Size: 20%'''
        self._add_paragraph(repro)
    
    def _add_section_image(self, section_key, index):
        """Add a single image from a section"""
        if section_key in self.images_dict:
            images = self.images_dict[section_key]
            if index < len(images):
                filename, caption = images[index]
                image_path = os.path.join(self.plots_dir, filename)
                add_image(self.document, image_path, caption)
